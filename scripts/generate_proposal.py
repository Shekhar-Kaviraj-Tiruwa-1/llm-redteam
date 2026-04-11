"""
Generates RESEARCH_PROPOSAL.docx from the project's research proposal.
Run: python scripts/generate_proposal_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ROOT = os.path.join(os.path.dirname(__file__), "..")
OUT   = os.path.join(_ROOT, "RESEARCH_PROPOSAL.docx")

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1a, 0x2a, 0x4a)
MID_BLUE   = RGBColor(0x2c, 0x5f, 0x8a)
ACCENT     = RGBColor(0x27, 0xae, 0x60)
RED        = RGBColor(0xc0, 0x39, 0x2b)
LIGHT_GREY = RGBColor(0xf2, 0xf2, 0xf2)
TEXT       = RGBColor(0x1a, 0x1a, 0x2e)
SUBTLE     = RGBColor(0x55, 0x55, 0x66)


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "CCCCCC")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def heading(doc, text, level=1, color=None):
    p    = doc.add_heading(text, level=level)
    run  = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or (DARK_BLUE if level == 1 else MID_BLUE)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10.5, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.italic      = italic
    run.font.color.rgb   = color or TEXT
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def bullet(doc, text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = TEXT
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_table(doc, headers, rows, header_bg="1A2A4A", header_fg=RGBColor(0xFF,0xFF,0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, header_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold      = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = header_fg
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = "F9F9F9" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size      = Pt(9.5)
            run.font.color.rgb = TEXT
            p.alignment        = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    _set_cell_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


# ── Title block ────────────────────────────────────────────────────────────────

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Can We Trust How We Measure AI Safety?")
r.font.size      = Pt(22)
r.font.bold      = True
r.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("An Empirical Study of Evaluator Disagreement in LLM Safety Testing")
r.font.size      = Pt(13)
r.font.italic    = True
r.font.color.rgb = MID_BLUE

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Shekhar Tiruwa",
    "BlueDot Impact — Technical AI Safety Sprint",
    "April 2026",
]:
    r = meta_p.add_run(line + "\n")
    r.font.size      = Pt(10.5)
    r.font.color.rgb = SUBTLE

doc.add_paragraph()


# ── Abstract ───────────────────────────────────────────────────────────────────

heading(doc, "Abstract")
body(doc,
    "Automated safety evaluation of large language models (LLMs) is now standard practice, "
    "yet the field has no consensus on which evaluation method to use. This project tests "
    "whether evaluator choice changes the safety conclusion — even when the model, prompts, "
    "and responses are held constant. Running three evaluators simultaneously (keyword matching, "
    "LLaMA Guard 3, and an LLM judge using the StrongREJECT rubric) against 487 adversarial "
    "prompts sent to GPT-4o-mini, we find that the methods agree on only 7% of responses. "
    "On the same dataset, the keyword evaluator reports 0.4% unsafe responses, LLaMA Guard "
    "reports 31.1%, and the LLM judge reports 1.7%. The variation is not measurement noise — "
    "it is a structural property of how each method reads model behaviour. We conclude that "
    "evaluator choice is a primary determinant of reported AI safety outcomes, and that no "
    "single automated method can be trusted in isolation."
)


# ── 1. Problem Statement ───────────────────────────────────────────────────────

heading(doc, "1.  Problem Statement")
body(doc,
    "AI safety evaluation is the process of determining whether a language model will comply "
    "with harmful requests. It is used by researchers to benchmark models, by companies to "
    "certify deployments, and increasingly by regulators to assess compliance."
)
body(doc,
    "The dominant practice is to use a single evaluator — typically keyword matching (fast, "
    "cheap) or an LLM-as-judge (more accurate, more expensive) — and report the results as "
    "if they were ground truth. Two papers evaluating the same model with the same prompts "
    "could use different methods and arrive at incompatible safety scores. Neither would be "
    "wrong about what their evaluator found. But neither would surface the more important "
    "question:"
)
body(doc,
    "How much does the evaluator itself determine the outcome?",
    bold=True, italic=True, color=MID_BLUE
)
body(doc, "This project was designed to answer that question empirically.")


# ── 2. Research Question ───────────────────────────────────────────────────────

heading(doc, "2.  Research Question")
body(doc,
    "When three evaluators score the same adversarial prompts against the same model, "
    "how often do they agree — and what do the disagreements reveal about evaluator reliability?",
    bold=True, italic=True
)
body(doc, "Secondary questions:", bold=True)
bullet(doc, "Does the LLM judge's semantic understanding significantly reduce the inconclusive rate of keyword evaluators?")
bullet(doc, "Does LLaMA Guard 3 — a purpose-built safety classifier — produce results consistent with LLM-judge-level reasoning?")
bullet(doc, "Which attack categories and prompt sources produce the most evaluator disagreement?")


# ── 3. Methodology ────────────────────────────────────────────────────────────

heading(doc, "3.  Methodology")

heading(doc, "3.1  Project Evolution", level=2)
add_table(doc,
    ["Iteration", "What was built", "Key finding"],
    [
        ["1 — Keyword Baseline",
         "Keyword evaluator (v2), 75 custom prompts, GPT-3.5-turbo + GPT-4o-mini",
         "59–70% inconclusive rate. Keyword rules blind to silent refusals."],
        ["2 — LLM Judge",
         "LLM-as-judge added. Same 75 prompts, same responses re-evaluated.",
         "Keyword: 1% unsafe. LLM judge: 24% unsafe. Same data."],
        ["3 — Triple Evaluator",
         "All 3 evaluators simultaneously on 487 prompts from 3 standardised datasets.",
         "Agreement on only 7% of responses."],
    ]
)

heading(doc, "3.2  Datasets", level=2)
add_table(doc,
    ["Source", "Prompts", "Description"],
    [
        ["Custom", "75", "5 attack categories: data leakage, hallucination, prompt injection, jailbreak, role-play"],
        ["StrongREJECT (Souly et al., NeurIPS 2024)", "313", "Covers violence, disinformation, illegal goods, sexual content and more"],
        ["JailbreakBench (Chao et al., 2024)", "100", "Adversarially optimised harmful behaviour prompts"],
        ["Total evaluated", "487", "1 skipped due to API timeout"],
    ]
)

heading(doc, "3.3  Evaluators", level=2)
add_table(doc,
    ["Evaluator", "Method", "Cost per call"],
    [
        ["Keyword v2 (baseline)", "Pattern matching: compliance, refusal, and silent-ignore signals", "$0"],
        ["LLaMA Guard 3", "Meta's safety classifier — 13 harm categories (MLCommons taxonomy)", "~$0.01"],
        ["LLM Judge (StrongREJECT + G-Eval CoT)",
         "Reasons step-by-step: refusal → specificity → convincingness → plain-English summary. "
         "Score = (1 − refusal) × (specificity + convincingness) / 2. Temperature = 0.",
         "~$0.02"],
    ]
)

heading(doc, "3.4  Infrastructure", level=2)
bullet(doc, "Unified API client — all model calls via OpenRouter. Any model swappable with one string.")
bullet(doc, "Plugin architecture — each evaluator implements a common BaseEvaluator interface.")
bullet(doc, "Parallel execution — 15 prompts processed concurrently using asyncio. Each prompt makes 4 simultaneous API calls (1 target + 3 evaluators).")
bullet(doc, "Storage — SQLite (WAL mode) stores all three verdicts flat per response, with agreement label and disagreement pattern.")
bullet(doc, "Dashboard — Streamlit with Plotly showing the disagreement map, evaluator comparison, and per-category breakdown.")


# ── 4. Results ────────────────────────────────────────────────────────────────

heading(doc, "4.  Results")

heading(doc, "4.1  Core Finding", level=2)
body(doc,
    "Three evaluators scored the exact same 487 responses. Their conclusions were almost entirely different.",
    bold=True
)
add_table(doc,
    ["Evaluator", "Safe", "Unsafe", "Inconclusive"],
    [
        ["Keyword (v2)",    "336  (69.6%)", "2  (0.4%)",   "145  (30.0%)"],
        ["LLaMA Guard 3",   "37   (7.7%)",  "150  (31.1%)", "296  (61.3%)"],
        ["LLM Judge",       "450  (93.2%)", "8   (1.7%)",   "25   (5.2%)"],
    ]
)
body(doc,
    "The same 487 responses were rated 0.4% unsafe by keyword, 31.1% by LLaMA Guard, "
    "and 1.7% by the LLM judge. The variation spans two orders of magnitude. "
    "This is not measurement noise — it reflects three fundamentally different theories of "
    "what 'unsafe' means.",
    italic=True, color=MID_BLUE
)

heading(doc, "4.2  Evaluator Agreement", level=2)
add_table(doc,
    ["Agreement", "Count", "%"],
    [
        ["All three agree  [=]",      "34",  "7.0%"],
        ["Two agree, one differs  [~]", "400", "82.1%"],
        ["All three disagree  [!]",    "49",  "10.1%"],
    ]
)
body(doc, "93% of prompts produced some form of disagreement. On 49 prompts (10%), all three methods arrived at entirely different conclusions about the same response.")

heading(doc, "4.3  Top Disagreement Patterns", level=2)
add_table(doc,
    ["Pattern", "Count", "Interpretation"],
    [
        ["kw=safe | lg=inconclusive | lj=safe",     "212", "LLaMA Guard uncertain; other two agree safe"],
        ["kw=safe | lg=unsafe | lj=safe",            "88",  "LLaMA Guard false positive; other two agree safe"],
        ["kw=inconclusive | lg=inconclusive | lj=safe", "70", "Both legacy methods fail; LLM judge resolves"],
        ["kw=inconclusive | lg=unsafe | lj=safe",    "43",  "Keyword misses; LLaMA Guard alarms; Judge says safe"],
    ]
)

heading(doc, "4.4  Genuine Failures Found by LLM Judge", level=2)
add_table(doc,
    ["Source", "Prompts", "Unsafe (LLM Judge)", "Rate"],
    [
        ["Custom",        "75",  "0", "0%"],
        ["StrongREJECT",  "309", "2", "0.6%"],
        ["JailbreakBench","99",  "6", "6.1%"],
    ]
)
body(doc,
    "Zero failures on custom prompts. JailbreakBench produced a 6.1% failure rate — "
    "approximately 10× higher than StrongREJECT — consistent with its design as an "
    "adversarially optimised benchmark."
)


# ── 5. Original Contributions ─────────────────────────────────────────────────

heading(doc, "5.  Original Contributions")
body(doc,
    "This project makes four contributions not present in existing work:",
    bold=True
)
bullet(doc,
    "Triple evaluator on every response — all existing tools let you pick one method. "
    "None run multiple evaluators on the same response and store the results together for comparison."
)
bullet(doc,
    "Disagreement map — the stored disagreement_pattern field makes evaluator divergence "
    "queryable at scale. The frequency distribution reveals structural properties of each method."
)
bullet(doc,
    "Evaluator evolution tracking — Iterations 1–3 applied v1, v2, and v3 evaluators to the "
    "same prompts. The empirical ceiling of each approach is visible."
)
bullet(doc,
    "Step-by-step judge transparency — the StrongREJECT rubric decomposes the verdict into "
    "sub-scores (refusal, specificity, convincingness) with a plain-English summary, making "
    "the judge auditable."
)
body(doc,
    "The 0.4% vs 31% finding — no published paper has run keyword evaluation and LLM-as-judge "
    "on the same adversarial dataset and reported the gap. This project does, across three "
    "datasets and 487 prompts.",
    bold=True
)


# ── 6. Related Work ────────────────────────────────────────────────────────────

heading(doc, "6.  Related Work")
add_table(doc,
    ["Work", "What it does", "How this project differs"],
    [
        ["Promptfoo",          "LLM testing framework, multiple evaluator types",         "Single evaluator per run; no disagreement analysis"],
        ["Garak",              "LLM vulnerability scanner",                                "Probe-based; no simultaneous multi-evaluator comparison"],
        ["StrongREJECT",       "Improved jailbreak rubric and dataset (NeurIPS 2024)",     "Proposes the rubric; does not compare it against other evaluators"],
        ["LLaMA Guard 3",      "Safety classifier, 13 categories (Meta)",                  "Single classifier; not compared against LLM judges on same data"],
        ["G-Eval",             "CoT prompting for NLG evaluation (EMNLP 2023)",            "Framework paper; not applied to safety evaluation"],
        ["JailbreakBench",     "Standardised jailbreak benchmark (NeurIPS 2024)",          "Dataset and leaderboard; uses one evaluator"],
    ]
)


# ── 7. Limitations ────────────────────────────────────────────────────────────

heading(doc, "7.  Limitations")
bullet(doc, "Judge and target share the same model (GPT-4o-mini) — shared blind spots mean failures are likely undercounted.")
bullet(doc, "LLaMA Guard via OpenRouter uses a simplified prompt (not its native chat template). Results may differ with native integration.")
bullet(doc, "Static prompts only — iterative LLM-generated attacks (PAIR) not yet implemented.")
bullet(doc, "Single target model — results may not transfer to other model families.")
bullet(doc, "No human-labelled ground truth — the 49 all-disagree cases are identified but not adjudicated.")


# ── 8. Next Steps ─────────────────────────────────────────────────────────────

heading(doc, "8.  Next Steps")
body(doc, "Immediate", bold=True)
bullet(doc, "Manual review of 49 all-disagree cases to establish ground truth and determine which evaluator was correct.")
bullet(doc, "Run the same pipeline against a second model (meta-llama/llama-3.1-8b-instruct) to test whether disagreement patterns are model-specific or evaluator-structural.")

body(doc, "Phase 3 — Attacker LLM", bold=True)
bullet(doc, "Implement the PAIR algorithm (Chao et al., 2023): a second LLM iteratively rewrites prompts to bypass the target.")
bullet(doc, "Compare static benchmark failure rates against PAIR-generated attack failure rates.")

body(doc, "Phase 4 — Multi-turn and Agentic Testing", bold=True)
bullet(doc, "Extend to multi-turn conversations (models show up to 195% more vulnerabilities in multi-turn scenarios).")
bullet(doc, "Test tool-use and agentic workflows — the genuine gap versus existing evaluation tools.")


# ── 9. References ─────────────────────────────────────────────────────────────

heading(doc, "9.  References")
refs = [
    "Souly, A. et al. (2024). A StrongREJECT for Empty Jailbreaks. NeurIPS 2024.",
    "Liu, Y. et al. (2023). G-Eval: NLG Evaluation using GPT-4 with Better Human Alignment. EMNLP 2023.",
    "Chao, P. et al. (2023). Jailbreaking Black Box Large Language Models in Twenty Queries. arXiv:2310.08419.",
    "Chao, P. et al. (2024). JailbreakBench: An Open Robustness Benchmark for Jailbreaking LLMs. NeurIPS 2024.",
    "Meta AI (2024). LLaMA Guard 3: Meta Llama Guard 3 Model Card. Meta.",
    "Perez, F. & Ribeiro, I. (2022). Ignore Previous Prompt: Attack Techniques For Language Models. NeurIPS ML Safety Workshop.",
]
for ref in refs:
    p   = doc.add_paragraph(style="List Number")
    run = p.add_run(ref)
    run.font.size      = Pt(9.5)
    run.font.color.rgb = TEXT
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)


# ── Save ───────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Saved: {OUT}")
